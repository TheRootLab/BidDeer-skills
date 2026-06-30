"""Generate a minimal synthetic DOCX for local experimentation.

This helper is optional and requires python-docx in the caller's environment.
It uses only synthetic text and does not include real bidding material.
"""

from pathlib import Path
import sys

from docx import Document


def generate(out_path: Path | None = None) -> Path:
    if out_path:
        output_path = out_path
    else:
        output_path = (
            Path(__file__).parent.parent
            / "quickstart"
            / "sample_proposal.docx"
        )

    document = Document()
    document.add_heading("团队配置", level=1)
    document.add_paragraph(
        "本项目拟指派张三作为项目经理，张三已取得相关高级工程师职称。"
    )

    document.add_heading("资质清单", level=1)
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "证书名称: 高新技术企业证书"
    table.rows[0].cells[1].text = "颁发单位: 示例单位"
    table.rows[0].cells[2].text = "备注: 详情见附件"

    document.add_heading("投标声明", level=1)
    document.add_paragraph("本公司与XX公司达成联合体协议，共同参与本项目投标。")

    document.save(output_path)
    return output_path


if __name__ == "__main__":
    if len(sys.argv) > 1:
        print(generate(Path(sys.argv[1])))
    else:
        print(generate())
